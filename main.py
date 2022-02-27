import sys
# Импорт библиотек для парсинга данных
import requests
from bs4 import BeautifulSoup
# Импорт библиотек для работы с Word
import os
from docx import Document
from docx.shared import Pt
from docx.shared import Mm
# Импорт библиотек для работы с консолью
import click
# Импорт библиотке для работы с json(файлом настройки)
import json


class Handler:
    heading = ""
    texts = []
    url = ""

    def __init__(self, url):
        self.url = url

    # Функция для парсинга
    def pars(self):
        self.heading = ""
        self.texts = []

        response = requests.get(self.url)
        soup = BeautifulSoup(response.text, 'lxml')

        # Нахождение текстовой информации
        setting = json.loads(open('setting.json', encoding="utf-8").read())
        if not setting.get('pars_teg'):
            click.echo("Некоректные настройки, укажите хотя бы 1 значение у ключа pars_teg")
            sys.exit()
        if not setting.get('ignore_class'):
            quotes = soup.find_all(setting.get('pars_teg'))
        else:
            quotes = soup.find_all(setting.get('pars_teg'), {"class": not setting.get('ignore_class')})
        self.heading = soup.find_all('h1')[0].text  # Нахождение заголовка

        for quote in quotes:
            # print(quote)
            if quote.find('a'):
                text_href = quote.find('a').text
                href = quote.find('a')['href']
                self.texts.append(quote.text.replace(text_href, text_href + f"[{href}]"))
            else:
                self.texts.append(quote.text)

    # Формирование word документа
    def doc(self):
        document = Document()
        style = document.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(16)
        run = document.add_paragraph().add_run(self.heading)
        run.font.size = Pt(24)
        run.bold = True
        document.add_paragraph(" ")
        for text in self.texts:
            paragraph = document.add_paragraph(text)
            fmt = paragraph.paragraph_format
            fmt.space_before = Mm(0)
            fmt.space_after = Mm(0)
            paragraph = document.add_paragraph(" ")
            fmt = paragraph.paragraph_format
            fmt.space_before = Mm(0)
            fmt.space_after = Mm(0)

        if "?" in self.url:
            self.url = self.url.replace("?", "/")

        path = self.url.split("://")[1].split("/")
        path = '\\'.join([str(x) for x in path])
        if not os.path.exists(path):
            os.makedirs(path)
        all_path = os.getcwd() + "\\" + path + '\\restyled.docx'
        document.save(all_path)
        return all_path


# Основная функция
@click.command()
@click.option('--command', prompt='Выберите цифру нужной команды', help='Возможные команды: 1 - отформатировать '
                                                                        'страницу сайта для комфортного чтения\n2 - '
                                                                        'Вывести настройки\n3 - Добавить новую '
                                                                        'инструкцию в настройки\n4 - Удалить '
                                                                        'инструкцию из настроек')
def cl_command(command):
    if command == '1':
        primary()
    elif command == '2':
        read_setting()
    elif command == '3':
        add_setting()
    elif command == '4':
        delete_setting()
    else:
        click.echo("Такой команды не существует, воспользуйтесь --help")


# Основная функция
@click.command()
@click.option('--url', prompt='Укажите ссылку на сайт', help='Ссылка на тот сайт, информацию откуда вы хотите взять')
def primary(url):
    handle = Handler(url)
    handle.pars()
    all_path = handle.doc()
    click.echo(f"файл создан по пути: {all_path}")


@click.command()
def read_setting():
    setting = json.loads(open('setting.json', encoding="utf-8").read())
    click.echo(setting)


@click.command()
@click.option('--name', prompt='Укажите имя ключа', help='Имя ключа(ignore_class, pars_teg)')
@click.option('--val', prompt='Укажите значение ключа', help='Значение ключа')
# Добавить новую инструкцию в файл настроек
def add_setting(name, val):
    setting = json.loads(open('setting.json', encoding="utf-8").read())
    if name in setting.keys():
        setting[name].append(val)
        with open('setting.json', "w") as file:
            json.dump(setting, file)
        click.echo("Найстройки изменены")
    else:
        click.echo("Такого ключа нет, попробуте pars_teg или ignore_class")


@click.command()
@click.option('--name', prompt='Укажите имя ключа', help='Имя ключа(ignore_class, pars_teg)')
@click.option('--val', prompt='Укажите значение ключа', help='Значение ключа')
# Удалить инструкцию из файла настроек
def delete_setting(name, val):
    setting = json.loads(open('setting.json', encoding="utf-8").read())
    if name in setting.keys():
        setting[name].remove(val)
        with open('setting.json', "w") as file:
            json.dump(setting, file)
        click.echo("Найстройки изменены")
    else:
        click.echo("Такого ключа нет, попробуте pars_teg или ignore_class")


if __name__ == '__main__':
    cl_command()
